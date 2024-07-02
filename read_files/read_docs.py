# reads large documents such as PDFs and .docx files.

# TODO pipelines for pdf, docx, etc...
# we want an analysis pipeline like the one for transcriptions
# different mqueue channel

# we can assume it'll just be one PDF, but there could be multiple. But it'll be a big PDF probably.

from os import PathLike
import PyPDF2
import docx
import httpx


def read_pdf_from_path(pdf_path: PathLike) -> str:
    """
    read a pdf from a path and return the text.
    """
    text = ""
    with open(pdf_path, "rb") as file:
        pdf = PyPDF2.PdfFileReader(file)
        for page in range(pdf.getNumPages()):
            text += pdf.getPage(page).extract_text()
    return text


def read_docx_from_path(docx_path: PathLike) -> str:
    """
    read a docx from a path and return the text.
    """
    text = ""
    with open(docx_path, "rb") as file:
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text
    return text


def read_pdf_from_url(pdf_url: str) -> str:
    """
    read a pdf from a URL and return the text.
    """
    text = ""
    with httpx.Client() as client:
        response = client.get(pdf_url)
        pdf = PyPDF2.PdfFileReader(response.content)
        for page in range(pdf.getNumPages()):
            text += pdf.getPage(page).extract_text()
    return text


def read_docx_from_url(docx_url: str) -> str:
    """
    read a docx from a URL and return the text.
    """
    text = ""
    with httpx.Client() as client:
        response = client.get(docx_url)
        doc = docx.Document(response.content)
        for para in doc.paragraphs:
            text += para.text
    return text